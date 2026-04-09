"""Generate a professional Word document for the Blend bug bounty report."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── Style definitions ────────────────────────────────────────────────────────

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    if level == 1:
        h.font.size = Pt(20)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(10)
    elif level == 2:
        h.font.size = Pt(14)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(12)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)


def add_colored_bg(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_font(cell, size=Pt(9), bold=False, color=None, font_name='Calibri'):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = size
            run.font.bold = bold
            run.font.name = font_name
            if color:
                run.font.color.rgb = color


def add_table(doc, headers, rows, col_widths=None, header_color="1A1A2E",
              alt_color="F5F5FA"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        add_colored_bg(cell, header_color)
        set_cell_font(cell, Pt(9), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            if r_idx % 2 == 1:
                add_colored_bg(cell, alt_color)
            set_cell_font(cell, Pt(9))
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, code, font_size=Pt(8.5)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F5" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = font_size
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def bold(p, text):
    run = p.add_run(text)
    run.bold = True
    return run


def mono(p, text):
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    return r


def add_severity_badge(doc, severity, color_hex):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  SEVERITY: {severity}  ")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rPr = run._r.get_or_add_rPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    rPr.append(shading)


def add_info_box(doc, text, bg_color="E8F0FE", border_color="1A73E8"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    pPr.append(shading)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_warning_box(doc, text):
    add_info_box(doc, text, bg_color="FFF3E0", border_color="E65100")


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
run = p.add_run("\u2501" * 60)
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(20)
run = p.add_run("BUG BOUNTY REPORT")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run("Blend Protocol v2 \u2014 Lending Pool Contract")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x66)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Utilization Rate Manipulation via Same-Asset Leverage Loops")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
run.font.italic = True

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
run = p.add_run("\u2501" * 60)
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
run.font.size = Pt(14)

# Metadata
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)

meta = [
    ["Submitted", "2026-03-14"],
    ["Severity", "CRITICAL"],
    ["Network", "Stellar Mainnet"],
    ["Affected Contract", "Lending Pool (submit / submit_with_allowance)"],
    ["Affected Pool", "Etherfuse (CDMAVJPF...HEQS5FPVAI)"],
    ["Affected Reserve", "USDC (CCW67TSZ...LEO7SJMI75)"],
    ["Status", "Confirmed on Mainnet \u2014 Positions Closed"],
]
table = doc.add_table(rows=len(meta), cols=2)
table.style = 'Table Grid'
for i, (k, v) in enumerate(meta):
    ck, cv = table.rows[i].cells[0], table.rows[i].cells[1]
    ck.text, cv.text = k, v
    add_colored_bg(ck, "1A1A2E")
    set_cell_font(ck, Pt(10), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    if v == "CRITICAL":
        add_colored_bg(cv, "FFCDD2")
        set_cell_font(cv, Pt(10), bold=True, color=RGBColor(0xB7, 0x1C, 0x1C))
    elif "Confirmed" in v:
        add_colored_bg(cv, "E8F5E9")
        set_cell_font(cv, Pt(10), bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
    else:
        set_cell_font(cv, Pt(10))
    ck.width, cv.width = Cm(4), Cm(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Summary', level=1)

p = doc.add_paragraph()
p.add_run(
    "Three findings are reported. Finding 1 has been demonstrated on mainnet. "
    "Findings 2 and 3 are theoretical / observational."
)

add_table(doc,
    ["#", "Title", "Severity", "Demonstrated"],
    [
        ["1", "Utilization rate manipulation via same-asset leverage loops", "Critical", "Yes \u2014 mainnet"],
        ["2", "Compounded oracle + utilization drain on thin-market assets", "High", "No \u2014 theoretical"],
        ["3", "TVL inflation via same-asset leverage loops", "Medium", "Yes \u2014 observable"],
    ],
    col_widths=[1, 8, 2.5, 3.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# FINDING 1
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Finding 1: Utilization Rate Manipulation', level=1)
add_severity_badge(doc, "CRITICAL", "B71C1C")

# --- Description ---
doc.add_heading('Description', level=2)

p = doc.add_paragraph()
p.add_run(
    "A same-asset leverage loop (supply USDC, borrow USDC, re-supply, re-borrow, repeated N "
    "times in a single "
)
mono(p, "submit_with_allowance")
p.add_run(
    " call) artificially dilutes a reserve's utilization rate. This allows a third party to "
    "withdraw liquidity that should be blocked by "
)
mono(p, "max_util")
p.add_run(
    ", and pushes borrowing rates into the R_3 penalty zone (~450% APR on Etherfuse USDC), "
    "causing direct financial loss to all borrowers in the reserve."
)

# --- Root Cause ---
doc.add_heading('Root Cause', level=2)

p = doc.add_paragraph()
p.add_run("The utilization cap (")
mono(p, "max_util")
p.add_run(") is enforced via ")
mono(p, "require_utilization_below_max")
p.add_run(", which is ")
bold(p, "only called from ")
mono(p, "apply_borrow")
p.add_run(
    ". Withdrawals are only checked against the hard 100% ceiling, not against the utilization cap."
)

add_table(doc,
    ["Operation", "util < 100%", "util < max_util (#1207)", "Health Check"],
    [
        ["Supply", "\u2014", "\u2014", "\u2014"],
        ["Withdraw", "Yes", "No", "\u2014"],
        ["SupplyCollateral", "\u2014", "\u2014", "\u2014"],
        ["WithdrawCollateral", "Yes", "No", "Yes"],
        ["Borrow", "Yes", "Yes", "Yes"],
        ["Repay", "\u2014", "\u2014", "\u2014"],
    ],
    col_widths=[4, 3.5, 4.5, 3],
)

add_info_box(doc,
    "During a same-asset loop, each borrow is individually checked against max_util. Because "
    "each iteration increases both supply and borrows, utilization converges toward c_factor. "
    "When c_factor \u2248 max_util, every borrow passes \u2014 even though the net effect is to "
    "mask the pool's true utilization."
)

# --- Impact ---
doc.add_heading('Impact', level=2)

# Impact 1
doc.add_heading('1. Interest Rate Manipulation \u2014 Loss of Funds', level=3)

p = doc.add_paragraph()
p.add_run("Pushing utilization above the target (80%) triggers the R_3 penalty slope:")

add_table(doc,
    ["Parameter", "Value"],
    [
        ["R_1 (base)", "5% APR"],
        ["R_2 (below target)", "20% APR"],
        ["R_3 (above target)", "500% APR"],
        ["Target utilization", "80%"],
        ["Utilization cap", "95%"],
    ],
    col_widths=[6, 6],
)

p = doc.add_paragraph()
p.add_run("At the post-exploit utilization of 97.21%:")

add_code_block(doc,
    "rate = 5% + 20% + ((0.97 - 0.80) / (1 - 0.80)) \u00d7 500% = ~450% APR",
    font_size=Pt(9.5),
)

p = doc.add_paragraph()
p.add_run("Health factor deterioration at 450% APR:")

add_table(doc,
    ["Starting HF", "Time to liquidation"],
    [
        ["1.02", "~1 day"],
        ["1.05", "~4 days"],
        ["1.10", "~8 days"],
    ],
    col_widths=[6, 6],
)

add_warning_box(doc,
    "The attacker profits directly: they hold a lending position (Wallet A) that earns the "
    "inflated APY. The loop/unloop (Wallet B) executes within a single Stellar ledger (~5 sec), "
    "so the attacker pays negligible borrowing interest. Only cost: gas fees.\n\n"
    "Within hours, users on the Blend Discord reported anomalous behavior on Etherfuse USDC."
)

# Impact 2-4
doc.add_heading('2. Liquidity Drain', level=3)

p = doc.add_paragraph()
p.add_run("Withdrawals during suppressed utilization remove real liquidity. Post-exploit:")

add_code_block(doc,
    "Total Supply:        82,011.50 USDC\n"
    "Total Borrows:       79,727.10 USDC\n"
    "Available Liquidity:  2,284.40 USDC  (2.8% of supply)\n"
    "Utilization:            97.21%\n"
    "Configured max_util:    95.00%",
    font_size=Pt(9.5),
)

doc.add_heading('3. Supplier Fund Lock', level=3)
p = doc.add_paragraph()
p.add_run(
    "Remaining suppliers cannot withdraw more than 2,284 USDC. They must wait for borrowers to repay."
)

doc.add_heading('4. Position Closure Deadlock', level=3)
p = doc.add_paragraph()
p.add_run("Looper positions may fail to unwind with ")
mono(p, "Error(Contract, #1207)")
p.add_run(" when the pool is above the utilization cap.")

# --- Reproduction Steps ---
doc.add_heading('Reproduction Steps', level=2)

p = doc.add_paragraph()
p.add_run("Two wallets are needed:")

p = doc.add_paragraph(style='List Bullet')
bold(p, "Wallet A")
p.add_run(" \u2014 existing lending position in the target pool")

p = doc.add_paragraph(style='List Bullet')
bold(p, "Wallet B")
p.add_run(" \u2014 will execute the leverage loop")

steps = [
    ("Deploy leverage loop (Wallet B):", " Build a submit_with_allowance transaction with N "
     "alternating supply_collateral + borrow requests for the same asset. Each borrow passes "
     "max_util individually because the preceding supply dilutes utilization."),
    ("Withdraw liquidity (Wallet A):", " While the loop is active, withdraw collateral. The "
     "withdrawal passes because apply_withdraw_collateral only checks util < 100%, not max_util."),
    ("Unwind the loop (Wallet B):", " Repay borrows and withdraw collateral. Utilization returns "
     "to its true level \u2014 now above max_util."),
    ("Collect interest (Wallet A):", " The lending position earns inflated APY (up to 500% APR). "
     "Steps 1\u20133 can be repeated."),
]

for i, (label, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{i}. ")
    r.font.bold = True
    bold(p, label)
    p.add_run(desc)

# --- On-Chain PoC ---
doc.add_heading('On-Chain Proof of Concept', level=2)

p = doc.add_paragraph()
p.add_run("Demonstrated on Stellar Mainnet, 2026-03-12/13.")

add_table(doc,
    ["Wallet", "Address", "Role"],
    [
        ["A (Lender)", "GBHD3V2XKX6DXHYZDSHA2\nUYZTO4MKB2R6QNSCDT4X\nEKNGTLPXT7A36EA", "Lending position"],
        ["B (Looper)", "GCR3VBVLYM5ZUBX63XMY\nBEY4EMAPVNCLORA4CWPA\n64CYEQQT53UCIQ36", "Leverage loops"],
    ],
    col_widths=[2.5, 7, 5.5],
)

p = doc.add_paragraph()
bold(p, "Leverage loops (Wallet B):")

add_table(doc,
    ["TX Hash (prefix)", "Timestamp", "Supply Added", "Borrows Added", "Leverage"],
    [
        ["78f090d376bb...", "2026-03-12 22:02", "2,000", "1,800", "10x"],
        ["b040e55d32aa...", "2026-03-12 22:36", "1,860", "1,660", "9.3x"],
        ["1f1ad7e2283a...", "2026-03-13 00:38", "1,860", "1,660", "9.3x"],
        ["a2d970953c6e...", "2026-03-13 22:49", "~10,000", "~9,000", "~10x"],
        ["485e928372ab...", "2026-03-13 22:52", "~102,000", "~91,800", "~10x"],
    ],
    col_widths=[3.5, 3.5, 2.5, 2.5, 2],
)

p = doc.add_paragraph()
bold(p, "Withdrawals during suppressed utilization (Wallet A):")

add_table(doc,
    ["TX Hash (prefix)", "Timestamp", "Amount (USDC)"],
    [
        ["c3b9fe7b...", "22:19:22", "318,131.16"],
        ["3cf3e120...", "22:20:43", "110,000.00"],
        ["f183e170...", "22:22:19", "27,000.00"],
        ["b2ee3423...", "22:23:11", "18,950.02"],
        ["8a2c6c84...", "22:25:05", "30,010.00"],
        ["eaf7bbbc...", "22:25:51", "25,106.92"],
        ["8f25cd9e...", "22:54:36", "5,614.15"],
    ],
    col_widths=[4, 4, 4],
)

p = doc.add_paragraph()
bold(p, "Loop unwinding (Wallet B):")

add_table(doc,
    ["TX Hash (prefix)", "Timestamp", "Event"],
    [
        ["62f8984f...", "2026-03-12 22:32", "repay 1,800 + withdraw 2,000"],
        ["7ebe93d7...", "2026-03-13 00:20", "repay 1,660 + withdraw 1,660"],
        ["b8997c65...", "2026-03-13 07:09", "withdraw 1,260"],
        ["b3668dfd...", "2026-03-13 23:07", "repay 81,600 + withdraw 91,800"],
    ],
    col_widths=[4, 4, 6],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.add_run("All positions are now closed. The pool remains at 97.21% utilization (above the 95% cap). ")
bold(p, "Video recording of the exploit is available upon request.")

# --- Suggested Fix ---
doc.add_heading('Suggested Fix', level=2)

p = doc.add_paragraph()
p.add_run("Add ")
mono(p, "do_check_max_util")
p.add_run(" to both ")
mono(p, "apply_withdraw")
p.add_run(" and ")
mono(p, "apply_withdraw_collateral")
p.add_run(":")

add_code_block(doc,
    "fn apply_withdraw_collateral(/* ... */) {\n"
    "    user.remove_collateral(e, &mut reserve, to_burn);\n"
    "    reserve.require_utilization_below_100(e);\n"
    "    actions.do_check_max_util(&reserve.asset);  // ADD THIS\n"
    "    actions.add_for_pool_transfer(&reserve.asset, tokens_out);\n"
    "    actions.do_check_health();\n"
    "}"
)

add_info_box(doc,
    "Trade-off: This also blocks legitimate withdrawals when utilization organically exceeds "
    "the cap. A more nuanced approach: only block withdrawals that would increase utilization "
    "above max_util.\n\n"
    "Additional defense-in-depth: enforce require_utilization_below_max for all reserves "
    "touched in validate_submit, not just those that were borrowed from."
)

doc.add_heading('Historical Precedent', level=2)

p = doc.add_paragraph()
p.add_run(
    "This attack pattern mirrors the "
)
bold(p, "Aave V2 / CRV incident (November 2022)")
p.add_run(
    ", where leverage loops manipulated utilization rates on the CRV reserve, resulting in "
)
bold(p, "~$1.6M in bad debt")
p.add_run(
    " absorbed by the Aave DAO. Aave subsequently implemented reserve-specific supply caps."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# FINDING 2
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Finding 2: Compounded Oracle + Utilization Attack', level=1)
add_severity_badge(doc, "HIGH", "C62828")

add_info_box(doc,
    "THEORETICAL: This attack has not been executed. It is included to illustrate the "
    "compounding risk when the utilization manipulation (Finding 1) interacts with the "
    "oracle architecture for thin-market Etherfuse assets."
)

# --- YieldBlox Background ---
doc.add_heading('Background: YieldBlox Exploit (Feb 22, 2026)', level=2)

p = doc.add_paragraph()
p.add_run(
    "The YieldBlox DAO pool was exploited for ~$10.8M via USTRY oracle manipulation. "
    "The YieldBlox pool used a "
)
bold(p, "raw Reflector oracle")
p.add_run(
    " with no circuit breakers \u2014 a $4 self-trade on SDEX inflated USTRY from ~$1.05 to "
    "~$106, and the oracle accepted it."
)

p = doc.add_paragraph()
p.add_run("Sources: Halborn, QuillAudits post-mortem analyses.")

# --- Etherfuse Oracle ---
doc.add_heading('Key Difference: Etherfuse Oracle Adaptor', level=2)

p = doc.add_paragraph()
p.add_run(
    "The Etherfuse pool uses a custom oracle adaptor with max_dev circuit breakers. "
    "The YieldBlox-style 100\u00d7 manipulation is "
)
bold(p, "not possible")
p.add_run(".")

add_table(doc,
    ["Asset", "max_dev", "Oracle", "Resolution"],
    [
        ["CETES", "5%", "Reflector (index 0)", "300s"],
        ["USTRY", "5%", "Reflector (index 0)", "300s"],
        ["TESOURO", "5%", "Reflector (index 0)", "300s"],
        ["USDC", "5%", "Reflector (index 1)", "300s"],
        ["XLM", "10%", "Reflector (index 1)", "300s"],
    ],
    col_widths=[3, 2, 5, 3],
)

# --- Why Still Exploitable ---
doc.add_heading('Why max_dev=5% Is Still Exploitable', level=2)

p = doc.add_paragraph()
bold(p, "5% = liquidation threshold. ")
p.add_run(
    "For CETES/TESOURO (c_factor=0.80), max leverage is ~5\u00d7, giving HF \u2248 1.05. "
    "A single 5% price drop liquidates these positions:"
)

add_code_block(doc, "HF 1.05 \u00d7 0.95 = 0.9975 \u2192 liquidatable", font_size=Pt(10))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
bold(p, "Price walking. ")
p.add_run(
    "If max_dev compares against the last accepted price, the price can be walked 5% "
    "every 300 seconds:"
)

add_table(doc,
    ["Time", "Cumulative drop", "Liquidates HF \u2264"],
    [
        ["5 min", "-5.0%", "1.05"],
        ["10 min", "-9.75%", "1.10"],
        ["30 min", "-26.5%", "1.36"],
        ["1 hour", "-46.0%", "Most positions"],
    ],
    col_widths=[3, 4, 5],
)

add_warning_box(doc,
    "SDEX depth for Etherfuse stablebonds is near-zero (~$1.39 bid depth for TESOURO). "
    "Moving the price 5% costs <$1 per update. Walking for 1 hour: ~$12."
)

# --- Compounded Scenario ---
doc.add_heading('Compounded Attack Scenario', level=2)

phases = [
    ("Phase 1 \u2014 Rate Spike:", " Use Finding 1 to spike USDC rates to ~450% APR, "
     "eroding borrower health factors."),
    ("Phase 2 \u2014 Oracle Manipulation:", " Push SDEX price of collateral asset down. "
     "max_dev=5% is sufficient for the first wave of liquidations."),
    ("Phase 3 \u2014 Compound:", " The attacks multiply: interest erosion + price drop."),
]
for i, (label, desc) in enumerate(phases, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    bold(p, label)
    p.add_run(desc)

add_table(doc,
    ["Starting HF", "After 1 day interest", "After 5% drop", "Combined", "Status"],
    [
        ["1.05", "1.037", "\u00d7 0.95 = 0.985", "0.985", "Liquidatable"],
        ["1.10", "1.087", "\u00d7 0.95 = 1.033", "1.033", "At risk"],
        ["1.05", "\u2014 (immed.)", "\u00d7 0.95 = 0.998", "0.998", "Liquidatable"],
    ],
    col_widths=[2.5, 3.5, 3, 2.5, 3.5],
)

# --- Why This Remains a Risk ---
doc.add_heading('Why This Remains a Risk', level=2)

risks = [
    ("max_dev=5% aligns with the liquidation threshold", " for max-leverage c_factor 0.80 positions."),
    ("Price walking", " may extend the attack beyond 5% over multiple oracle periods."),
    ("SDEX liquidity is near-zero", " for Etherfuse stablebonds \u2014 manipulation costs are negligible."),
    ("No minimum volume check", " \u2014 a single self-trade on an empty order book is accepted."),
    ("The utilization manipulation compounds the effect", " \u2014 rate erosion weakens health factors before the oracle push."),
]
for i, (label, desc) in enumerate(risks, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"{i}. ")
    r.font.bold = True
    bold(p, label)
    p.add_run(desc)

# --- Mitigations ---
doc.add_heading('Suggested Mitigations', level=2)

mitigations = [
    ("Fixed reference for max_dev:", " Compare against TWAP/median, not last accepted price."),
    ("Minimum volume thresholds:", " Reject prices from periods with negligible SDEX volume."),
    ("Multi-source oracle:", " Require agreement between Reflector and a secondary source."),
    ("Tighter max_dev:", " Reduce to 2\u20133% for assets with c_factor \u2265 0.80."),
    ("Collateral value caps:", " Limit max collateral for thin-market assets."),
]
for label, desc in mitigations:
    p = doc.add_paragraph(style='List Bullet')
    bold(p, label)
    p.add_run(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# FINDING 3
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Finding 3: TVL Inflation via Same-Asset Leverage Loops', level=1)
add_severity_badge(doc, "MEDIUM", "E65100")

# --- Description ---
doc.add_heading('Description', level=2)

p = doc.add_paragraph()
p.add_run(
    "Same-asset leverage loops artificially inflate a pool's reported Total Value Locked (TVL). "
    "Each loop iteration adds both supply and borrows, causing on-chain "
)
mono(p, "total_supply")
p.add_run(
    " and "
)
mono(p, "total_borrows")
p.add_run(
    " to grow geometrically while the user's actual economic exposure remains constant. "
    "This misrepresents the pool's real size to depositors, analytics platforms, and risk models."
)

# --- Root Cause ---
doc.add_heading('Root Cause', level=2)

p = doc.add_paragraph()
p.add_run(
    "Blend reports TVL as the sum of all bToken balances (supply positions). There is no mechanism "
    "to distinguish between real deposits backed by external capital and synthetic deposits created "
    "by re-supplying borrowed funds. A same-asset loop creates supply and borrow entries that net "
    "to zero economic exposure but inflate both sides of the balance sheet."
)

# --- Impact ---
doc.add_heading('Impact', level=2)

doc.add_heading('Inflation Factor', level=3)

p = doc.add_paragraph()
p.add_run("For collateral factor ")
bold(p, "c")
p.add_run(", a single user depositing ")
bold(p, "X")
p.add_run(" can inflate apparent TVL by:")

add_code_block(doc, "TVL inflation factor = 1 / (1 - c)", font_size=Pt(10))

add_table(doc,
    ["Asset", "c_factor", "Max TVL Inflation", "$200 deposit appears as"],
    [
        ["USDC", "0.95", "20\u00d7", "$4,000"],
        ["USTRY", "0.90", "10\u00d7", "$2,000"],
        ["CETES", "0.80", "5\u00d7", "$1,000"],
    ],
    col_widths=[3, 2.5, 3.5, 4],
)

doc.add_heading('Observed On-Chain', level=3)

p = doc.add_paragraph()
p.add_run("The leverage loop transactions from Finding 1 demonstrate this directly:")

add_table(doc,
    ["TX Hash (prefix)", "Real Deposit", "Apparent Supply", "Apparent Borrows", "Inflation"],
    [
        ["78f090d3...", "200 USDC", "~2,000 USDC", "~1,800 USDC", "10\u00d7"],
        ["a2d97095...", "~1,000 USDC", "~10,000 USDC", "~9,000 USDC", "~10\u00d7"],
        ["485e9283...", "~10,200 USDC", "~102,000 USDC", "~91,800 USDC", "~10\u00d7"],
    ],
    col_widths=[3, 2.5, 3, 3, 2.5],
)

add_warning_box(doc,
    "A single user with ~11,400 USDC of real capital created ~114,000 USDC of apparent "
    "supply \u2014 inflating the pool's reported TVL by ~100,000 USDC."
)

doc.add_heading('Downstream Consequences', level=3)

consequences = [
    ("Depositor deception:", " New depositors perceive the pool as larger and more liquid than "
     "it actually is, leading to misallocation of capital."),
    ("Risk model corruption:", " Automated risk systems using TVL as input overestimate the "
     "pool's safety margin."),
    ("Analytics distortion:", " DeFi aggregators (DeFiLlama, etc.) report inflated TVL, "
     "misrepresenting protocol adoption and capital base."),
    ("Governance manipulation:", " If BLND incentives or governance weight are tied to pool TVL, "
     "leverage loops can farm rewards or influence votes disproportionately."),
]
for label, desc in consequences:
    p = doc.add_paragraph(style='List Bullet')
    bold(p, label)
    p.add_run(desc)

# --- Reproduction Steps ---
doc.add_heading('Reproduction Steps', level=2)

repro_steps = [
    ("Deposit X USDC", " into the pool."),
    ("Execute a submit_with_allowance", " with N alternating supply_collateral + borrow "
     "requests for USDC (same steps as Finding 1)."),
    ("Observe", " that the pool's reported total_supply has increased by ~X / (1 - c) "
     "while the user's net economic position is unchanged."),
    ("The inflated TVL persists", " for as long as the loop remains open."),
]
for i, (label, desc) in enumerate(repro_steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"{i}. ")
    r.font.bold = True
    bold(p, label)
    p.add_run(desc)

# --- Suggested Fix ---
doc.add_heading('Suggested Fix', level=2)

fixes = [
    ("Compute effective TVL:", " Net out self-referential positions: "
     "effective_supply = total_supply - same_asset_borrows."),
    ("Per-reserve supply caps:", " Limit the maximum inflation any single actor can create."),
    ("Display net TVL:", " Show total_supply - total_borrows on front-ends and analytics "
     "for a more accurate picture of real capital."),
]
for label, desc in fixes:
    p = doc.add_paragraph(style='List Bullet')
    bold(p, label)
    p.add_run(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDICES
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Appendix A: Affected Code', level=1)

p = doc.add_paragraph()
mono(p, "apply_borrow")
p.add_run(" \u2014 the only place ")
mono(p, "do_check_max_util")
p.add_run(" is called:")

add_code_block(doc,
    "fn apply_borrow(/* ... */) {\n"
    "    user.add_liabilities(e, &mut reserve, d_tokens_minted);\n"
    "    reserve.require_utilization_below_100(e);\n"
    "    actions.do_check_max_util(&reserve.asset);  // <-- ONLY here\n"
    "    actions.add_for_pool_transfer(&reserve.asset, request.amount);\n"
    "    actions.do_check_health();\n"
    "}"
)

p = doc.add_paragraph()
mono(p, "apply_withdraw_collateral")
p.add_run(" \u2014 missing max_util check:")

add_code_block(doc,
    "fn apply_withdraw_collateral(/* ... */) {\n"
    "    user.remove_collateral(e, &mut reserve, to_burn);\n"
    "    reserve.require_utilization_below_100(e);  // only hard cap\n"
    "    actions.add_for_pool_transfer(&reserve.asset, tokens_out);\n"
    "    actions.do_check_health();\n"
    "    // MISSING: actions.do_check_max_util(&reserve.asset);\n"
    "}"
)

doc.add_heading('Appendix B: Utilization Dilution Math', level=1)

add_code_block(doc,
    "Added supply  = X \u00d7 (1 - c^(N+1)) / (1 - c)\n"
    "Added borrows = X \u00d7 c \u00d7 (1 - c^N) / (1 - c)\n\n"
    "Ratio added_borrows / added_supply \u2192 c  as N \u2192 \u221e\n\n"
    "If pool utilization U > c, the loop pulls U down toward c.\n"
    "When c \u2248 max_util, this creates the exploitation window.",
    font_size=Pt(9),
)

doc.add_heading('Appendix C: Interest Rate Model', level=1)

add_code_block(doc,
    "If utilization \u2264 target:\n"
    "    rate = R_1 + (utilization / target) \u00d7 R_2\n\n"
    "If utilization > target:\n"
    "    rate = R_1 + R_2 + ((utilization - target) / (1 - target)) \u00d7 R_3\n\n"
    "Rate Modifier (RM) amplifies/dampens reactively.\n\n"
    "Etherfuse USDC: R_1=5%, R_2=20%, R_3=500%, target=80%",
    font_size=Pt(9),
)

# Footer
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
run = p.add_run("\u2501" * 60)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "This report was prepared based on on-chain transaction analysis and source code review "
    "of blend-contracts-v2. All transaction hashes reference Stellar Mainnet. The exploit was "
    "conducted for security research purposes and all positions have been closed."
)
r.font.size = Pt(8.5)
r.font.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ─────────────────────────────────────────────────────────────────────

out_path = os.path.join(os.path.dirname(__file__), "Blend_Bug_Bounty_Report.docx")
doc.save(out_path)
print(f"Report saved to: {out_path}")
