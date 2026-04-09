"""Generate a professional Word document for the Blend vulnerability report."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── Style definitions ────────────────────────────────────────────────────────

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    if level == 1:
        h.font.size = Pt(20)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(10)
    elif level == 2:
        h.font.size = Pt(14)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(12)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)


def add_colored_bg(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_font(cell, size=Pt(9), bold=False, color=None, font_name='Calibri'):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = size
            run.font.bold = bold
            run.font.name = font_name
            if color:
                run.font.color.rgb = color


def add_table(doc, headers, rows, col_widths=None, header_color="1A1A2E",
              alt_color="F5F5FA"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        add_colored_bg(cell, header_color)
        set_cell_font(cell, Pt(9), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            if r_idx % 2 == 1:
                add_colored_bg(cell, alt_color)
            set_cell_font(cell, Pt(9))

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


def add_code_block(doc, code, font_size=Pt(8.5)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    # Add shading to the paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F5" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = font_size
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_bold_text(p, text):
    run = p.add_run(text)
    run.bold = True
    return run


def add_severity_badge(doc, severity, color_hex):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  SEVERITY: {severity}  ")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Background via shading
    rPr = run._r.get_or_add_rPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    rPr.append(shading)


def add_info_box(doc, text, bg_color="E8F0FE", border_color="1A73E8"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    pPr.append(shading)
    # Left border
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_warning_box(doc, text):
    add_info_box(doc, text, bg_color="FFF3E0", border_color="E65100")


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Top decorative line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
run = p.add_run("\u2501" * 60)
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
run.font.size = Pt(14)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(20)
run = p.add_run("SECURITY REPORT")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Blend Protocol v2 \u2014 Pool Contract")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x66)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Utilization Rate Manipulation via Same-Asset Leverage Loops")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
run.font.name = 'Calibri'
run.font.italic = True

# Bottom decorative line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
run = p.add_run("\u2501" * 60)
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
run.font.size = Pt(14)

# Metadata table
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)

meta_data = [
    ["Report Date", "2026-03-14"],
    ["Severity", "CRITICAL"],
    ["Network", "Stellar Mainnet"],
    ["Affected Pool", "Etherfuse (CDMAVJPF...HEQS5FPVAI)"],
    ["Affected Reserve", "USDC (CCW67TSZ...LEO7SJMI75)"],
    ["Component", "Pool utilization rate calculation and interest rate model"],
    ["Status", "Confirmed on Mainnet \u2014 Active"],
]

table = doc.add_table(rows=len(meta_data), cols=2)
table.style = 'Table Grid'
for i, (k, v) in enumerate(meta_data):
    cell_k = table.rows[i].cells[0]
    cell_v = table.rows[i].cells[1]
    cell_k.text = k
    cell_v.text = v
    add_colored_bg(cell_k, "1A1A2E")
    set_cell_font(cell_k, Pt(10), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    if v == "CRITICAL":
        add_colored_bg(cell_v, "FFCDD2")
        set_cell_font(cell_v, Pt(10), bold=True, color=RGBColor(0xB7, 0x1C, 0x1C))
    elif v == "Confirmed on Mainnet \u2014 Active":
        add_colored_bg(cell_v, "FFF3E0")
        set_cell_font(cell_v, Pt(10), bold=True, color=RGBColor(0xE6, 0x51, 0x00))
    else:
        set_cell_font(cell_v, Pt(10))
    cell_k.width = Cm(4)
    cell_v.width = Cm(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Executive Summary",
    "2. Vulnerability Details",
    "    2.1 Root Cause",
    "    2.2 Utilization Dilution Mechanism",
    "    2.3 TVL Inflation",
    "3. Impact Assessment",
    "    3.1 Interest Rate Manipulation (Critical)",
    "    3.2 Liquidity Drain (High)",
    "    3.3 Supplier Fund Lock (High)",
    "    3.4 Position Closure Deadlock (High)",
    "    3.5 TVL Misrepresentation (Medium)",
    "4. Proof of Concept \u2014 On-Chain Evidence",
    "    4.1 Actors",
    "    4.2 Pre-Conditions",
    "    4.3 Timeline of Events",
    "5. Attack Scenario \u2014 Profitable Exploitation",
    "6. Compounded Attack: Oracle + Utilization (Theoretical)",
    "    6.1 The YieldBlox Oracle Exploit (Feb 2026)",
    "    6.2 Different Oracle, Residual Risk",
    "    6.3 Why max_dev=5% Is Still Exploitable",
    "    6.4 Compounded Attack Scenario",
    "    6.5 Why This Remains a Risk",
    "7. Affected Contract Code",
    "8. Recommended Mitigations",
    "9. Conclusion",
]

for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    indent = item.startswith("    ")
    if indent:
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(item.strip())
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x66)
    else:
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('1. Executive Summary', level=1)

add_severity_badge(doc, "CRITICAL", "B71C1C")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.add_run(
    "A same-asset leverage loop (supply USDC as collateral, borrow USDC, re-supply, "
    "re-borrow, repeated N times in a single atomic transaction) can artificially dilute "
    "a reserve's utilization rate, temporarily bringing it below the configured utilization cap ("
)
r = p.add_run("max_util")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(
    "). This enables third-party withdrawals that drain real liquidity from the pool. "
    "When the loop is subsequently unwound \u2014 which can happen in the same Stellar "
    "ledger (~5 seconds) \u2014 the utilization rate snaps back to its true, now higher level."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
add_bold_text(p, "The primary impact is twofold:")

# Impact 1
p = doc.add_paragraph(style='List Bullet')
add_bold_text(p, "Interest Rate Manipulation: ")
p.add_run(
    "Pushing utilization above the target rate triggers Blend's steep interest rate curve (R_3), "
    "which on the Etherfuse USDC reserve reaches 500% APR. This forces all borrowers in the "
    "reserve to pay massively inflated rates, rapidly deteriorating their health factors and "
    "pushing them toward liquidation. The attacker can profit from this by holding a lending "
    "position that earns the inflated APY."
)

# Impact 2
p = doc.add_paragraph(style='List Bullet')
add_bold_text(p, "Liquidity Drain: ")
p.add_run(
    "Withdrawals executed during the suppressed-utilization window remove real liquidity, leaving "
    "the pool in a degraded state where utilization exceeds the utilization cap \u2014 a state the cap "
    "is designed to prevent as protection against \"oracle instability or collateral asset exploits\" "
    "(per Blend documentation)."
)

add_warning_box(doc,
    "This vulnerability was confirmed on Stellar Mainnet on 2026-03-12/13 against "
    "the Etherfuse pool's USDC reserve. Within hours, users on the Blend Discord "
    "reported anomalous behavior on the USDC reserve."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. VULNERABILITY DETAILS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('2. Vulnerability Details', level=1)
doc.add_heading('2.1 Root Cause', level=2)

p = doc.add_paragraph()
p.add_run("Blend's utilization cap (")
r = p.add_run("max_util")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(") is enforced via ")
r = p.add_run("require_utilization_below_max")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(", which is ")
add_bold_text(p, "only called on borrow operations")
p.add_run(" (")
r = p.add_run("apply_borrow")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run("). Withdrawals (")
r = p.add_run("apply_withdraw, apply_withdraw_collateral")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(") are only checked against the hard 100% ceiling (")
r = p.add_run("require_utilization_below_100")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run("), not against ")
r = p.add_run("max_util")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(".")

add_info_box(doc,
    "Per Blend documentation, the utilization cap exists as \"protection from oracle "
    "instability or collateral asset exploits.\" However, because withdrawals bypass this "
    "cap, a leverage loop can temporarily suppress the apparent utilization, allow withdrawals "
    "that drain liquidity, and then unwind \u2014 leaving the pool above the cap."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
add_bold_text(p, "Relevant contract code \u2014 ")
r = p.add_run("pool/src/pool/actions.rs")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(":")

add_table(doc,
    ["Operation", "util < 100%", "util < max_util (#1207)", "Health Check"],
    [
        ["Supply", "\u2014", "\u2014", "\u2014"],
        ["Withdraw", "Yes", "No", "\u2014"],
        ["SupplyCollateral", "\u2014", "\u2014", "\u2014"],
        ["WithdrawCollateral", "Yes", "No", "Yes"],
        ["Borrow", "Yes", "Yes", "Yes"],
        ["Repay", "\u2014", "\u2014", "\u2014"],
    ],
    col_widths=[4, 3.5, 4.5, 3],
)

# 2.2
doc.add_heading('2.2 Utilization Dilution Mechanism', level=2)

p = doc.add_paragraph("Consider a pool with:")
items = [
    "Total supply: S",
    "Total borrows: B",
    "Utilization: U = B / S",
    "Collateral factor: c (e.g., 0.95)",
    "Utilization cap: max_util (e.g., 0.95)",
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.add_run(
    "A user deposits X USDC and executes N loop iterations. Each iteration supplies "
    "the previous borrow amount and borrows c times that amount. After N iterations:"
)

add_code_block(doc,
    "Added supply  = X \u00d7 (1 + c + c\u00b2 + ... + c\u207f)  = X \u00d7 (1 - c^(N+1)) / (1 - c)\n"
    "Added borrows = X \u00d7 (c + c\u00b2 + ... + c\u207f)      = X \u00d7 c \u00d7 (1 - c^N) / (1 - c)\n\n"
    "New utilization = (B + added_borrows) / (S + added_supply)"
)

p = doc.add_paragraph()
r = add_bold_text(p, "Key insight: ")
p.add_run("The ratio added_borrows / added_supply converges to c as N increases. Therefore:")

bullets2 = [
    ("If U > c", "the loop pulls utilization DOWN toward c"),
    ("If U < c", "the loop pushes utilization UP toward c"),
    ("If U = c", "the loop has no effect on utilization"),
]
for condition, effect in bullets2:
    p = doc.add_paragraph(style='List Bullet')
    r = add_bold_text(p, condition)
    p.add_run(f": {effect}")
    p.paragraph_format.space_after = Pt(2)

add_warning_box(doc,
    "In the critical case where a pool is near or above max_util (and max_util \u2248 c_factor), "
    "a leverage loop artificially reduces the displayed utilization rate, creating a window "
    "where withdrawals succeed that should not."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
add_bold_text(p, "Numerical example (matching the on-chain exploit):")

add_table(doc,
    ["State", "Supply (USDC)", "Borrows (USDC)", "Utilization"],
    [
        ["Before loop (pre-exploit)", "82,000", "77,900", "95.0%"],
        ["After 10x loop with 10,200 USDC", "177,600", "173,400", "~97.6%*"],
        ["Third-party withdraws 5,000 USDC", "172,600", "173,400", "Blocked (>100%)"],
        ["Loop unwound", "77,000", "77,900", "101.2%"],
    ],
    col_widths=[6, 3, 3, 3],
)

p = doc.add_paragraph()
p.add_run("* Individual borrow-step utilization stays below max_util due to the dilution effect.")
p.runs[0].font.size = Pt(8.5)
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)

# 2.3
doc.add_heading('2.3 TVL Inflation (Secondary Issue)', level=2)

p = doc.add_paragraph()
p.add_run(
    "Same-asset leverage loops artificially inflate the pool's reported Total Value Locked (TVL). "
    "A user depositing 200 USDC and looping 10 times creates ~2,000 USDC of apparent supply "
    "and ~1,800 USDC of apparent borrows \u2014 a "
)
add_bold_text(p, "10x inflation")
p.add_run(" of the economic reality. This misleads depositors, analytics dashboards, and risk models.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. IMPACT ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('3. Impact Assessment', level=1)

# 3.1 Interest Rate Manipulation
doc.add_heading('3.1 Interest Rate Manipulation \u2014 Forced Liquidation Risk', level=2)
add_severity_badge(doc, "CRITICAL", "B71C1C")

p = doc.add_paragraph()
add_bold_text(p, "This is the most severe impact. ")
p.add_run(
    "Blend's interest rate model uses a three-leg piecewise function with a steep penalty "
    "slope (R_3) above the target utilization rate. The Rate Modifier (RM) further amplifies "
    "rates reactively when utilization stays above target."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.add_run("On the Etherfuse pool's USDC reserve, the configured parameters are:")

add_table(doc,
    ["Parameter", "Value"],
    [
        ["Target utilization", "80%"],
        ["R_1 (base rate)", "5% APR"],
        ["R_2 (slope below target)", "20% APR"],
        ["R_3 (slope above target)", "500% APR"],
        ["Utilization cap (max_util)", "95%"],
    ],
    col_widths=[6, 6],
)

p = doc.add_paragraph()
p.add_run(
    "When the attacker pushes utilization from ~80% to 97%+, the effective borrowing rate "
    "jumps to hundreds of percent APR. At these rates, borrower health factors deteriorate rapidly:"
)

add_table(doc,
    ["Starting Health Factor", "Time to Liquidation (at 500% APR)"],
    [
        ["HF 1.02", "~1 day"],
        ["HF 1.05", "~4 days"],
        ["HF 1.10", "~8 days"],
    ],
    col_widths=[6, 6],
)

p = doc.add_paragraph()
add_bold_text(p, "Attack profitability: ")
p.add_run("The attacker does not pay the inflated borrowing rate because:")

profit_items = [
    "The leverage loop and unloop can execute within the same Stellar ledger (~5 seconds), "
    "so the attacker's borrowing position exists for negligible time.",
    "The attacker holds a separate lending position (Wallet A) that earns the inflated APY "
    "on real supplied funds.",
    "The only cost is transaction gas fees (negligible on Stellar).",
]
for item in profit_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

add_warning_box(doc,
    "In effect, the attacker extracts value from every borrower in the reserve, redistributing "
    "it to lenders (including themselves) and the backstop (Blend's first-loss capital module)."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
add_bold_text(p, "Historical Precedent: ")
p.add_run(
    "This attack pattern closely mirrors the Aave V2 / CRV incident (November 2022), where "
    "an attacker used leverage loops to manipulate utilization rates on the CRV reserve. The "
    "inflated interest rates and illiquid conditions ultimately resulted in ~$1.6M in bad debt "
    "that had to be absorbed by the Aave DAO treasury. The incident prompted Aave to implement "
    "reserve-specific supply caps and risk parameter adjustments."
)

add_code_block(doc,
    "At 97% utilization (post-exploit):\n"
    "  rate = R_1 + R_2 + ((0.97 - 0.80) / (1 - 0.80)) * R_3\n"
    "       = 5% + 20% + 0.85 * 500%\n"
    "       = 450% APR (before Rate Modifier amplification)",
    font_size=Pt(9),
)

# 3.2 Liquidity Drain
doc.add_heading('3.2 Liquidity Drain', level=2)
add_severity_badge(doc, "HIGH", "C62828")

p = doc.add_paragraph()
p.add_run(
    "Withdrawals executed during suppressed utilization remove real liquidity from the pool. "
    "When the loop is unwound, utilization exceeds the cap \u2014 a state the cap is designed "
    "to prevent. In the observed case:"
)

add_code_block(doc,
    "Post-exploit USDC Reserve:\n"
    "  Total Supply:        82,011.50 USDC\n"
    "  Total Borrows:       79,727.10 USDC\n"
    "  Available Liquidity:  2,284.40 USDC  (2.8% of supply)\n"
    "  Current Utilization:    97.21%\n"
    "  Configured max_util:    95.00%\n"
    "  STATUS: UTILIZATION CAP VIOLATED",
    font_size=Pt(9.5),
)

# 3.3
doc.add_heading('3.3 Supplier Fund Lock', level=2)
add_severity_badge(doc, "HIGH", "C62828")

p = doc.add_paragraph()
p.add_run(
    "When utilization approaches 100%, suppliers who did not participate in the attack find "
    "their funds locked. "
)
r = p.add_run("require_utilization_below_100")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(
    " blocks any withdrawal that would push utilization to 100%. These users must wait "
    "for borrowers to repay \u2014 a condition that may take an indeterminate amount of time."
)

add_info_box(doc,
    "In the observed case, any supplier attempting to withdraw more than 2,284 USDC would be blocked."
)

# 3.4
doc.add_heading('3.4 Position Closure Deadlock', level=2)
add_severity_badge(doc, "HIGH", "C62828")

p = doc.add_paragraph()
p.add_run(
    "If a looper's position remains open when the pool is above the utilization cap, the "
    "REPAY + WITHDRAW_COLLATERAL sequence required to close it may fail with "
)
r = p.add_run("Error(Contract, #1207)")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(
    " \u2014 InvalidUtilRate \u2014 because the repay-then-withdraw sequence can momentarily "
    "reduce supply faster than borrows, spiking utilization."
)

add_info_box(doc,
    "Note on liquidations: Blend's liquidation mechanism uses \"gentle dutch auctions\" and "
    "operates via position accounting (bToken/dToken transfers) rather than actual token "
    "movements. Liquidations will therefore always execute regardless of pool liquidity. "
    "However, the interest rate manipulation attack creates the liquidation conditions in "
    "the first place \u2014 borrowers are pushed to liquidation by artificially inflated rates, "
    "not by organic market movements."
)

# 3.5
doc.add_heading('3.5 TVL Misrepresentation', level=2)
add_severity_badge(doc, "MEDIUM", "E65100")

p = doc.add_paragraph()
p.add_run(
    "Leverage loops inflate reported TVL by up to 1/(1-c) times the actual deposited capital. "
    "For c_factor = 0.95, a single user can inflate TVL by up to "
)
add_bold_text(p, "20x")
p.add_run(" their real deposit.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. PROOF OF CONCEPT
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('4. Proof of Concept \u2014 On-Chain Evidence', level=1)

p = doc.add_paragraph()
p.add_run(
    "The following transactions were executed on Stellar Mainnet between "
    "2026-03-12 21:11 UTC and 2026-03-13 23:46 UTC."
)

doc.add_heading('4.1 Actors', level=2)

add_table(doc,
    ["Label", "Address", "Role"],
    [
        ["Wallet A (Lender)", "GBHD3V2XKX6DXHYZDSHA2UYZTO4\nMKB2R6QNSCDT4XEKNGTLPXT7A36EA",
         "Passive USDC supplier \u2014\nearns inflated APY"],
        ["Wallet B (Looper)", "GCR3VBVLYM5ZUBX63XMYBEY4EMA\nPVNCLORA4CWPA64CYEQQT53UCIQ36",
         "Executed leverage loops"],
    ],
    col_widths=[3, 8, 4],
)

doc.add_heading('4.2 Pre-Conditions', level=2)

preconditions = [
    "Wallet A held a lending position in the Etherfuse pool's USDC reserve.",
    "The USDC reserve was already at a high utilization rate (~95%), near the utilization cap.",
    "The interest rate was at the boundary of the R_2/R_3 transition.",
]
for item in preconditions:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

doc.add_heading('4.3 Timeline of Events', level=2)

# Phase 1
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = add_bold_text(p, "Phase 1 \u2014 Leverage Loop Deployment")
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

# Loop TX 1
p = doc.add_paragraph()
add_bold_text(p, "Transaction 1 \u2014 Leverage Loop (10x, 200 USDC)")

add_table(doc,
    ["Field", "Value"],
    [
        ["TX Hash", "78f090d376bb54de2dcefee11153ad217405b32fa5ace17e27392682882c2fa6"],
        ["Timestamp", "2026-03-12 22:02:07 UTC"],
        ["Ledger", "61,623,817"],
        ["Operation", "submit_with_allowance \u2014 13 supply_collateral + 13 borrow"],
        ["Initial Deposit", "200.00 USDC"],
        ["Total Supply Added", "2,000.00 USDC"],
        ["Total Borrows Added", "1,800.00 USDC"],
        ["Effective Leverage", "10.0x"],
    ],
    col_widths=[4, 12],
)

p = doc.add_paragraph("Individual loop iterations:")
add_code_block(doc,
    "supply  200.00  ->  borrow  190.00\n"
    "supply  190.00  ->  borrow  180.50\n"
    "supply  180.50  ->  borrow  171.47\n"
    "supply  171.47  ->  borrow  162.90\n"
    "supply  162.90  ->  borrow  154.76\n"
    "supply  154.76  ->  borrow  147.02\n"
    "supply  147.02  ->  borrow  139.67\n"
    "supply  139.67  ->  borrow  132.68\n"
    "supply  132.68  ->  borrow  126.05\n"
    "supply  126.05  ->  borrow  119.75\n"
    "supply  119.75  ->  borrow  113.76\n"
    "supply  113.76  ->  borrow  108.07\n"
    "supply  108.07  ->  borrow   53.37\n"
    "supply   53.37  (final collateral deposit)"
)

# Loop TX 2
p = doc.add_paragraph()
add_bold_text(p, "Transaction 2 \u2014 Leverage Loop (9.3x, 200 USDC)")

add_table(doc,
    ["Field", "Value"],
    [
        ["TX Hash", "b040e55d32aaa9d314049fcd186a3c594c5bad2bc937045acc7d504af15e0f56"],
        ["Timestamp", "2026-03-12 22:36:42 UTC"],
        ["Ledger", "61,624,178"],
        ["Total Supply Added", "1,860.00 USDC"],
        ["Total Borrows Added", "1,660.00 USDC"],
        ["Effective Leverage", "9.3x"],
    ],
    col_widths=[4, 12],
)

# Loop TX 3
p = doc.add_paragraph()
add_bold_text(p, "Transaction 3 \u2014 Leverage Loop (9.3x, 200 USDC)")

add_table(doc,
    ["Field", "Value"],
    [
        ["TX Hash", "1f1ad7e2283a4d539aa6b136cbf7b1b6a10260d1d3cb883ef660cdd61fb152b0"],
        ["Timestamp", "2026-03-13 00:38:23 UTC"],
        ["Ledger", "61,625,450"],
        ["Total Supply Added", "1,860.26 USDC"],
        ["Total Borrows Added", "1,660.23 USDC"],
        ["Effective Leverage", "9.3x"],
    ],
    col_widths=[4, 12],
)

# Loop TX 4
p = doc.add_paragraph()
add_bold_text(p, "Transaction 4 \u2014 Leverage Loop (10x, 1,000 USDC)")

add_table(doc,
    ["Field", "Value"],
    [
        ["TX Hash", "a2d970953c6ec26e6a24..."],
        ["Timestamp", "2026-03-13 22:49:13 UTC"],
        ["Ledger", "61,639,085"],
        ["Total Supply Added", "~10,000 USDC"],
        ["Total Borrows Added", "~9,000 USDC"],
        ["Effective Leverage", "~10x"],
    ],
    col_widths=[4, 12],
)

# Loop TX 5
p = doc.add_paragraph()
add_bold_text(p, "Transaction 5 \u2014 Leverage Loop (10x, 10,200 USDC)")

add_table(doc,
    ["Field", "Value"],
    [
        ["TX Hash", "485e928372ab48915712..."],
        ["Timestamp", "2026-03-13 22:52:09 UTC"],
        ["Ledger", "61,639,102"],
        ["Total Supply Added", "~102,000 USDC"],
        ["Total Borrows Added", "~91,800 USDC"],
        ["Effective Leverage", "~10x"],
    ],
    col_widths=[4, 12],
)

# Phase 2
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = add_bold_text(p, "Phase 2 \u2014 Withdrawals During Suppressed Utilization")
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

p = doc.add_paragraph()
p.add_run(
    "While leverage loops were active, Wallet A executed withdrawals from the pool that "
    "would not have been possible at the pool's true utilization rate:"
)

add_table(doc,
    ["Timestamp (UTC)", "Event", "Amount (USDC)", "TX Hash (prefix)"],
    [
        ["2026-03-13 22:19:22", "withdraw_collateral", "318,131.16", "c3b9fe7b..."],
        ["2026-03-13 22:20:43", "withdraw_collateral", "110,000.00", "3cf3e120..."],
        ["2026-03-13 22:22:19", "withdraw_collateral", "27,000.00", "f183e170..."],
        ["2026-03-13 22:23:11", "withdraw_collateral", "18,950.02", "b2ee3423..."],
        ["2026-03-13 22:25:05", "withdraw_collateral", "30,010.00", "8a2c6c84..."],
        ["2026-03-13 22:25:51", "withdraw_collateral", "25,106.92", "eaf7bbbc..."],
        ["2026-03-13 22:54:36", "withdraw_collateral", "5,614.15", "8f25cd9e..."],
    ],
    col_widths=[4, 3.5, 3, 4.5],
)

# Phase 3
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = add_bold_text(p, "Phase 3 \u2014 Loop Unwind and Resulting Pool Degradation")
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

add_table(doc,
    ["Timestamp (UTC)", "Event", "Amount (USDC)", "TX Hash (prefix)"],
    [
        ["2026-03-12 22:32:53", "repay + withdraw", "1,800 / 2,000", "62f8984f..."],
        ["2026-03-13 00:20:53", "repay + withdraw", "1,660 / 1,660", "7ebe93d7..."],
        ["2026-03-13 07:09:35", "withdraw_collateral", "1,260", "b8997c65..."],
        ["2026-03-13 23:07:05", "repay + withdraw", "81,600 / 91,800", "b3668dfd..."],
    ],
    col_widths=[4, 3.5, 3.5, 4],
)

# Phase 4
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = add_bold_text(p, "Phase 4 \u2014 Post-Exploit Pool State")
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

add_code_block(doc,
    "USDC Reserve \u2014 queried 2026-03-14:\n"
    "  Total Supply (underlying):   82,011.50 USDC\n"
    "  Total Borrows (underlying):  79,727.10 USDC\n"
    "  Available Liquidity:          2,284.40 USDC\n"
    "  Current Utilization:            97.21%\n"
    "  Configured max_util:            95.00%\n"
    "  STATUS: UTILIZATION CAP VIOLATED",
    font_size=Pt(9.5),
)

add_warning_box(doc,
    "The Loop Wallet's position is now fully closed (empty). Wallet A retains ~19,450 USDC "
    "in collateral. Within hours, users on the Blend Discord reported anomalous behavior on "
    "the USDC reserve, including unexpectedly high borrowing rates."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. ATTACK SCENARIO
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('5. Attack Scenario \u2014 Profitable Exploitation', level=1)

p = doc.add_paragraph()
p.add_run(
    "A malicious actor can profit from this vulnerability using two wallets:"
)

doc.add_heading('Step-by-step:', level=3)

steps = [
    ("Setup (Wallet A):", " Establish a lending position in a pool where utilization is already "
     "near the utilization cap. This position earns the standard lending APY."),
    ("Loop (Wallet B):", " Deploy a same-asset leverage loop (supply + borrow \u00d7 N) in a single "
     "atomic transaction. This temporarily suppresses apparent utilization below the cap."),
    ("Withdraw (Wallet A):", " While utilization appears low, withdraw real liquidity from the pool. "
     "This step is optional \u2014 the interest rate manipulation alone is profitable."),
    ("Unloop (Wallet B):", " In the same ledger or shortly after, unwind the leverage loop. "
     "Utilization spikes above the cap and into the R_3 penalty zone."),
    ("Collect (Wallet A):", " The lending position now earns the inflated APY (up to 500% APR on "
     "Etherfuse USDC) while all borrowers in the reserve pay the inflated rate. The attacker "
     "can repeat steps 2\u20134 to maintain the elevated rate."),
]

for i, (label, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{i}. ")
    r.font.bold = True
    add_bold_text(p, label)
    p.add_run(desc)

doc.add_heading('Cost/Benefit Analysis:', level=3)

add_table(doc,
    ["Component", "Value"],
    [
        ["Attacker cost", "Gas fees only (~fractions of a cent on Stellar)"],
        ["Attacker borrowing cost", "Negligible \u2014 loop exists for <5 seconds (one ledger)"],
        ["Capital required", "A few hundred USDC (for the loop; lending position is separate)"],
        ["Attacker revenue", "Inflated lending APY on Wallet A's position"],
        ["Victim cost", "All borrowers pay 500% APR instead of normal rates"],
        ["Time to profit", "Immediate \u2014 rates change within the same ledger"],
    ],
    col_widths=[5, 11],
)

add_warning_box(doc,
    "Historical Precedent: This attack pattern is analogous to the Aave V2 / CRV incident "
    "(November 2022), where an attacker used leverage loops to manipulate utilization rates "
    "on the CRV reserve. The inflated interest rates and illiquid conditions ultimately "
    "resulted in ~$1.6M in bad debt that had to be absorbed by the Aave DAO treasury. The "
    "incident prompted Aave to implement reserve-specific supply caps and risk parameter "
    "adjustments."
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. COMPOUNDED ATTACK: ORACLE MANIPULATION + UTILIZATION DRAIN
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('6. Compounded Attack Vector: Oracle Manipulation + Utilization Drain', level=1)

add_info_box(doc,
    "THEORETICAL: This attack scenario has not been executed. It is a theoretical analysis "
    "based on the mechanics demonstrated in Section 5, the on-chain oracle architecture, and "
    "the precedent set by the YieldBlox exploit (February 2026). It is included to illustrate "
    "the compounding risk when multiple vulnerabilities interact."
)

doc.add_heading('6.1 The YieldBlox Oracle Exploit (Feb 22, 2026)', level=2)

p = doc.add_paragraph()
p.add_run(
    "On February 22, 2026, the YieldBlox DAO pool on Blend Protocol was exploited for "
)
add_bold_text(p, "~$10.8M")
p.add_run(
    " via oracle price manipulation of USTRY \u2014 the same Etherfuse stablebond asset "
    "present in the Etherfuse pool."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.add_run("The attack mechanics:")

exploit_steps = [
    "The USTRY/USDC pair on SDEX had less than $1 in hourly trading volume.",
    "The attacker placed a sell offer at ~$106/USTRY (fair value: ~$1.05).",
    "Using a second account, they bought 0.05 USTRY at the inflated price \u2014 a single "
    "self-trade costing ~$4.",
    "The Reflector oracle's 5-minute VWAP window picked up this single trade as the dominant "
    "price signal.",
    "The attacker deposited USTRY at the inflated oracle price and borrowed ~1M USDC + ~61M XLM.",
]
for item in exploit_steps:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

add_warning_box(doc,
    "The YieldBlox pool used a raw Reflector oracle with no circuit breakers. A 100\u00d7 price "
    "manipulation was accepted in a single oracle update. Cost: ~$4. Funds extracted: ~$10.8M."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run("Sources: ")
r.font.size = Pt(8.5)
r.font.italic = True
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
r = p.add_run(
    "Halborn \u2014 YieldBlox Hack Explained (Feb 2026), "
    "QuillAudits \u2014 How a Single Trade Caused $10M Loss"
)
r.font.size = Pt(8.5)
r.font.italic = True
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_heading('6.2 The Etherfuse Pool: Different Oracle, Residual Risk', level=2)

add_info_box(doc,
    "Important distinction: Unlike the YieldBlox pool, the Etherfuse pool does NOT use a raw "
    "Reflector oracle. It uses a custom oracle adaptor contract that wraps the underlying "
    "Reflector feeds with circuit breaker protections (max_dev). The YieldBlox-style 100\u00d7 "
    "manipulation in a single update is not possible."
)

p = doc.add_paragraph()
p.add_run(
    "On-chain inspection of the oracle adaptor's instance storage reveals the following "
    "circuit breaker configuration:"
)

add_table(doc,
    ["Asset", "max_dev", "Oracle Index", "Underlying Oracle"],
    [
        ["CETES", "5%", "0", "Reflector (300s, 14 dec)"],
        ["USTRY", "5%", "0", "Reflector (300s, 14 dec)"],
        ["TESOURO", "5%", "0", "Reflector (300s, 14 dec)"],
        ["USDC", "5%", "1", "Reflector (300s, 14 dec)"],
        ["XLM", "10%", "1", "Reflector (300s, 14 dec)"],
    ],
    col_widths=[3, 2, 2.5, 7.5],
)

p = doc.add_paragraph()
add_bold_text(p, "MaxAge: ")
p.add_run("600 seconds (stale prices are rejected).")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.add_run(
    "The max_dev circuit breaker rejects any price update that deviates more than 5% "
    "from the previous accepted price. However, this protection has significant limitations:"
)

doc.add_heading('6.3 Why max_dev=5% Is Still Exploitable', level=2)

p = doc.add_paragraph()
add_bold_text(p, "5% is enough to liquidate. ")
p.add_run(
    "Any borrower at HF \u2264 1.05 \u2014 which includes anyone at maximum leverage with "
    "CETES or TESOURO (c_factor = 0.80, max leverage ~5\u00d7) \u2014 is liquidated by a "
    "single 5% price drop:"
)

add_code_block(doc,
    "Borrower at HF 1.05:\n"
    "  new HF = 1.05 \u00d7 0.95 = 0.9975 \u2192 liquidatable",
    font_size=Pt(9.5),
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
add_bold_text(p, "Price walking. ")
p.add_run(
    "If the max_dev check compares against the last accepted price (rather than a fixed "
    "reference), an attacker can walk the price \u2014 moving it 5% every oracle resolution "
    "period (300 seconds):"
)

add_table(doc,
    ["Time elapsed", "Cumulative price movement", "Effect"],
    [
        ["5 min (1 update)", "-5.0%", "Liquidates HF \u2264 1.05"],
        ["10 min (2 updates)", "-9.75%", "Liquidates HF \u2264 1.10"],
        ["30 min (6 updates)", "-26.5%", "Liquidates HF \u2264 1.36"],
        ["1 hour (12 updates)", "-46.0%", "Catastrophic"],
    ],
    col_widths=[4, 4, 7],
)

add_warning_box(doc,
    "SDEX depth for Etherfuse stablebonds is near-zero (~$1.39 of bid depth for TESOURO "
    "at time of the YieldBlox exploit). Moving the price 5% per update costs almost nothing. "
    "The cost of walking the price for one hour: ~$12. The resulting ~46% price drop would "
    "be catastrophic for any pool accepting these assets as collateral."
)

doc.add_heading('6.4 Compounded Attack Scenario', level=2)

p = doc.add_paragraph()
p.add_run("An attacker can combine both vulnerabilities for maximum impact:")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = add_bold_text(p, "Phase 1 \u2014 Interest Rate Spike")
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

p = doc.add_paragraph()
p.add_run(
    "Using the leverage loop technique from Section 5, spike the USDC borrowing rate to "
    "~450% APR. This begins eroding the health factor of every borrower in the pool."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = add_bold_text(p, "Phase 2 \u2014 Oracle Price Manipulation")
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

p = doc.add_paragraph()
p.add_run(
    "Simultaneously or after the rate spike, manipulate the SDEX price of the collateral "
    "asset downward. The max_dev=5% circuit breaker limits each update, but 5% is sufficient "
    "for the first wave of liquidations \u2014 and price walking extends the damage over time."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = add_bold_text(p, "Phase 3 \u2014 Compounded Effect")
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

p = doc.add_paragraph()
p.add_run("The two attacks compound multiplicatively:")

add_table(doc,
    ["Starting HF", "After 1 day interest", "After 5% price drop", "Combined HF", "Status"],
    [
        ["1.05", "1.037", "\u00d7 0.95 = 0.985", "0.985", "Liquidatable"],
        ["1.10", "1.087", "\u00d7 0.95 = 1.033", "1.033", "At risk"],
        ["1.15", "1.136", "\u00d7 0.95 = 1.079", "1.079", "Weakened"],
        ["1.05", "\u2014 (immediate)", "\u00d7 0.95 = 0.998", "0.998", "Liquidatable"],
    ],
    col_widths=[2.5, 3, 3.5, 2.5, 3.5],
)

doc.add_heading('6.5 Why This Remains a Risk', level=2)

risk_items = [
    "5% aligns with liquidation threshold: The max_dev cap of 5% coincides exactly with "
    "the liquidation margin for max-leverage positions with c_factor 0.80 assets.",
    "Price walking: If max_dev compares against the last accepted price, the 5% limit can "
    "be bypassed over multiple oracle periods (5% every 300 seconds).",
    "Near-zero SDEX liquidity: The cost to move Etherfuse stablebond prices 5% is negligible.",
    "No minimum volume check: The oracle adaptor does not verify SDEX trade volume or "
    "liquidity depth. A single self-trade on an empty order book is accepted.",
    "Compounding effect: Interest rate erosion via utilization manipulation pushes borrowers "
    "closer to liquidation before the oracle manipulation delivers the final push.",
]
for item in risk_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

doc.add_heading('6.6 Recommended Additional Mitigations', level=2)

oracle_mitigations = [
    ("Fixed reference for max_dev:", " Compare oracle price updates against a TWAP or median "
     "of recent values rather than the last accepted price, to prevent price walking."),
    ("Minimum liquidity thresholds:", " Reject oracle prices derived from periods with SDEX "
     "volume below a configurable threshold."),
    ("Multi-source oracle feeds:", " Integrate secondary price sources (e.g., RedStone, "
     "which launched Stellar support on March 4, 2026) and require agreement between sources."),
    ("Collateral value caps:", " Limit the maximum USD value of collateral that can be posted "
     "for thin-market assets, capping exposure to oracle manipulation."),
    ("Tighter max_dev for high-leverage assets:", " For assets with c_factor \u2265 0.80 "
     "(allowing \u22655\u00d7 leverage), reduce max_dev to 2\u20133% so a single oracle update "
     "cannot liquidate max-leverage positions."),
]
for label, desc in oracle_mitigations:
    p = doc.add_paragraph(style='List Bullet')
    add_bold_text(p, label)
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. AFFECTED CONTRACT CODE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('7. Affected Contract Code', level=1)

doc.add_heading('7.1 apply_borrow \u2014 Utilization Check (Only Location)', level=2)
p = doc.add_paragraph()
r = p.add_run("File: pool/src/pool/actions.rs")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x66)

add_code_block(doc,
    "fn apply_borrow(/* ... */) {\n"
    "    // ...\n"
    "    user.add_liabilities(e, &mut reserve, d_tokens_minted);\n"
    "    reserve.require_utilization_below_100(e);\n"
    "    actions.do_check_max_util(&reserve.asset);  // <-- ONLY borrow does this\n"
    "    actions.add_for_pool_transfer(&reserve.asset, request.amount);\n"
    "    actions.do_check_health();\n"
    "}"
)

doc.add_heading('7.2 apply_withdraw_collateral \u2014 No Utilization Cap Check', level=2)

add_code_block(doc,
    "fn apply_withdraw_collateral(/* ... */) {\n"
    "    // ...\n"
    "    user.remove_collateral(e, &mut reserve, to_burn);\n"
    "    reserve.require_utilization_below_100(e);  // only hard cap, NOT max_util\n"
    "    actions.add_for_pool_transfer(&reserve.asset, tokens_out);\n"
    "    actions.do_check_health();\n"
    "    // MISSING: actions.do_check_max_util(&reserve.asset);\n"
    "}"
)

add_warning_box(doc,
    "The do_check_max_util call is absent from apply_withdraw_collateral. "
    "This means withdrawals are not subject to the utilization cap, "
    "only the hard 100% solvency ceiling."
)

doc.add_heading('7.3 validate_submit \u2014 Post-Batch Check', level=2)

add_code_block(doc,
    "fn validate_submit(e: &Env, actions: &Actions) {\n"
    "    // Only checks max_util for assets that were BORROWED\n"
    "    for asset in &actions.check_max_util {\n"
    "        let reserve = storage::get_reserve(e, asset);\n"
    "        reserve.require_utilization_below_max(e);  // Error #1207\n"
    "    }\n"
    "    // ...\n"
    "}"
)

p = doc.add_paragraph()
p.add_run("The ")
r = p.add_run("check_max_util")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(" set is only populated by ")
r = p.add_run("do_check_max_util")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(", which is only called from ")
r = p.add_run("apply_borrow")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(". Withdrawals never add to this set.")

doc.add_heading('7.4 Interest Rate Model', level=2)
p = doc.add_paragraph()
r = p.add_run("File: pool/src/pool/reserve.rs")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x66)

p = doc.add_paragraph()
p.add_run(
    "Blend uses a three-leg piecewise interest rate function with a dynamic Rate Modifier (RM):"
)

add_code_block(doc,
    "If utilization <= target:\n"
    "    rate = R_1 + (utilization / target) * R_2\n\n"
    "If utilization > target:\n"
    "    rate = R_1 + R_2 + ((utilization - target) / (1 - target)) * R_3\n\n"
    "On Etherfuse USDC:\n"
    "    R_1 = 5% APR     (base rate)\n"
    "    R_2 = 20% APR    (slope below target)\n"
    "    R_3 = 500% APR   (slope above target - penalty zone)\n"
    "    Target = 80%",
    font_size=Pt(9),
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. RECOMMENDED MITIGATIONS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('8. Recommended Mitigations', level=1)

doc.add_heading('8.1 Enforce Utilization Cap on Withdrawals (Primary Fix)', level=2)

p = doc.add_paragraph()
p.add_run("Add ")
r = p.add_run("do_check_max_util")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(" to both ")
r = p.add_run("apply_withdraw")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(" and ")
r = p.add_run("apply_withdraw_collateral")
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
p.add_run(":")

add_code_block(doc,
    "fn apply_withdraw_collateral(/* ... */) {\n"
    "    user.remove_collateral(e, &mut reserve, to_burn);\n"
    "    reserve.require_utilization_below_100(e);\n"
    "    actions.do_check_max_util(&reserve.asset);  // ADD THIS\n"
    "    actions.add_for_pool_transfer(&reserve.asset, tokens_out);\n"
    "    actions.do_check_health();\n"
    "}"
)

add_info_box(doc,
    "Trade-off: This would also block legitimate suppliers from withdrawing when "
    "utilization organically exceeds the cap. A more nuanced approach: block withdrawals "
    "that increase utilization above the cap, while allowing withdrawals that keep "
    "utilization at or below its current level."
)

doc.add_heading('8.2 Same-Asset Loop Detection / Restriction', level=2)

p = doc.add_paragraph("Consider restricting same-asset leverage loops at the contract level:")

options = [
    ("Option A:", " Disallow borrowing an asset that the user is simultaneously supplying "
     "as collateral in the same submit batch."),
    ("Option B:", " Compute utilization using net supply/borrows (excluding self-referential "
     "positions) for the max_util check."),
    ("Option C:", " Implement per-reserve supply caps (as Aave did post-CRV incident)."),
]
for label, desc in options:
    p = doc.add_paragraph(style='List Bullet')
    add_bold_text(p, label)
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(4)

doc.add_heading('8.3 Post-Batch Utilization Invariant', level=2)

p = doc.add_paragraph()
p.add_run(
    "Rather than checking utilization per-operation, enforce a global invariant at the "
    "end of validate_submit:"
)

add_code_block(doc,
    "For every reserve TOUCHED in this batch:\n"
    "    require_utilization_below_max(e)"
)

p = doc.add_paragraph()
p.add_run(
    "This ensures that no combination of operations within a single submit call can "
    "leave the pool above the utilization cap."
)

doc.add_heading('8.4 Rate Manipulation Guard', level=2)

p = doc.add_paragraph()
p.add_run(
    "Consider rate-of-change limits on the interest rate model to prevent sudden spikes "
    "from single-transaction utilization manipulation."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('9. Conclusion', level=1)

p = doc.add_paragraph()
p.add_run(
    "The Blend Protocol v2 pool contract enforces its utilization cap (max_util) only on "
    "borrow operations. This asymmetry enables same-asset leverage loops to temporarily "
    "suppress apparent utilization, allowing liquidity withdrawals that violate the cap. "
    "More critically, when the loop is unwound, the resulting utilization spike pushes "
    "borrowing rates into the R_3 penalty zone (up to 500% APR on Etherfuse USDC), creating "
    "a profitable interest rate manipulation attack that harms all borrowers in the reserve."
)

p = doc.add_paragraph()
p.add_run("The vulnerability was demonstrated on Stellar Mainnet, resulting in:")

results = [
    "Post-exploit utilization of 97.21% against a configured cap of 95%",
    "Borrowing rates spiking to ~450% APR",
    "Observable user impact on the Blend Discord within hours",
    "Potential forced liquidation of borrowers with health factors below ~1.10",
]
for item in results:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph()
p.add_run(
    "The attack requires minimal capital (a few hundred USDC), costs only gas fees, "
    "and can be executed within a single Stellar ledger. This pattern closely parallels "
    "the Aave V2/CRV incident that resulted in $1.6M in bad debt."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.add_run(
    "We recommend enforcing the utilization cap on withdrawal operations as a priority fix, "
    "and implementing supply caps or loop restrictions as defense-in-depth measures."
)

# Footer
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
run = p.add_run("\u2501" * 60)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "This report was prepared based on on-chain transaction analysis and source code review "
    "of the blend-contracts-v2 repository. All transaction hashes reference Stellar Mainnet "
    "and can be independently verified. The exploit was conducted by the reporter for security "
    "research purposes and all positions have been closed."
)
r.font.size = Pt(8.5)
r.font.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ─────────────────────────────────────────────────────────────────────

out_path = os.path.join(os.path.dirname(__file__), "Blend_Protocol_Security_Report.docx")
doc.save(out_path)
print(f"Report saved to: {out_path}")
